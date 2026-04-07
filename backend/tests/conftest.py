from collections.abc import Generator
from pathlib import Path

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.middleware.auth import AuthenticatedUser, get_current_user

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session")
def generated_pdf_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    pdf_path = tmp_path_factory.mktemp("generated-fixtures") / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "CS 101 — Spring 2025 Syllabus\n"
        "Homework 1 due January 30, 2025\n"
        "Midterm Exam on March 10, 2025\n"
        "Final Project due April 20, 2025",
    )
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture(scope="session")
def generated_docx_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    docx_path = tmp_path_factory.mktemp("generated-fixtures") / "sample.docx"
    doc = Document()
    doc.add_paragraph("Midterm Exam — March 10, 2025")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Quiz 1"
    table.rows[0].cells[1].text = "Feb 14, 2025"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture()
def test_user() -> AuthenticatedUser:
    return AuthenticatedUser(
        id="00000000-0000-0000-0000-000000000001",
        email="test@example.com",
        access_token="test-token",
    )


@pytest.fixture()
def api_client() -> Generator[TestClient, None, None]:
    with TestClient(app, raise_server_exceptions=False) as client:
        yield client


@pytest.fixture()
def authenticated_client(
    test_user: AuthenticatedUser,
) -> Generator[TestClient, None, None]:
    def override_current_user() -> AuthenticatedUser:
        return test_user

    app.dependency_overrides[get_current_user] = override_current_user
    try:
        with TestClient(app, raise_server_exceptions=False) as client:
            yield client
    finally:
        app.dependency_overrides.clear()
