"""Text extraction from PDF, Word, and screenshot images."""

import asyncio
import base64
import io
from collections.abc import Awaitable, Callable
from zipfile import BadZipFile, ZipFile

import filetype
import fitz
import pdfplumber
from docx import Document
from fastapi import HTTPException, UploadFile
from fastapi.concurrency import run_in_threadpool
from openai import AsyncOpenAI
from openai.types.chat import ChatCompletionContentPartParam

from app.config import settings
from app.services import extraction_cache

MAX_PDF_PAGES = 100
PDF_PROCESSING_TIMEOUT_SECONDS = 30
# Cap a rasterized PDF page at ~25 megapixels so a maliciously huge MediaBox
# cannot exhaust memory during vision rendering (page.get_pixmap).
MAX_RENDER_PIXELS = 25_000_000
# Reject DOCX whose total uncompressed size exceeds this, to stop zip bombs
# (python-docx would otherwise decompress the whole archive into memory).
MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
# Explicit timeout (seconds) for OpenAI vision calls. The SDK default is 600s;
# bound worker exposure to a hung request. Vision OCR over several high-detail
# pages is slower than a text completion, so allow more headroom than the text
# extraction path.
OPENAI_VISION_TIMEOUT_SECONDS = 120.0
# Hard cap on screenshots per request, enforced BEFORE any billable vision
# call. Vision token cost scales with the number of images, so this is the
# primary cost guard for the screenshot path.
MAX_SCREENSHOT_IMAGES = 10
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"
ALLOWED_IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp"}
ALLOWED_DOC_MIMES = {PDF_MIME, DOCX_MIME}


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes via PyMuPDF, falling back to pdfplumber."""
    text = ""
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            if len(doc) > MAX_PDF_PAGES:
                raise HTTPException(
                    status_code=413,
                    detail=(
                        f"PDF exceeds {MAX_PDF_PAGES}-page limit. "
                        "Please split it into smaller files."
                    ),
                )
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                page_text = page.get_text()
                if isinstance(page_text, str):
                    text += page_text
    except HTTPException:
        raise
    except Exception:
        pass

    if len(text.strip()) < 50:
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                if len(pdf.pages) > MAX_PDF_PAGES:
                    raise HTTPException(
                        status_code=413,
                        detail=(
                            f"PDF exceeds {MAX_PDF_PAGES}-page limit. "
                            "Please split it into smaller files."
                        ),
                    )
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except HTTPException:
            raise
        except Exception:
            pass

    return text.strip()


def _pdf_pages_to_base64_images(
    data: bytes, max_pages: int | None = None, dpi: int | None = None
) -> list[str]:
    """Convert PDF pages to base64-encoded PNG images for vision API."""
    if max_pages is None:
        max_pages = settings.vision_max_pages
    if dpi is None:
        dpi = settings.vision_dpi
    images: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i in range(min(len(doc), max_pages)):
            page = doc[i]
            zoom = dpi / 72
            # Clamp the render so a maliciously huge MediaBox cannot blow up
            # memory: cap the rasterized page at MAX_RENDER_PIXELS.
            rect = page.rect
            est_px = (rect.width * zoom) * (rect.height * zoom)
            if est_px > MAX_RENDER_PIXELS:
                zoom *= (MAX_RENDER_PIXELS / est_px) ** 0.5
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            images.append(base64.b64encode(img_bytes).decode("utf-8"))
    return images


async def _extract_pdf_via_vision(data: bytes) -> str:
    """Use OpenAI vision to OCR scanned PDF pages."""
    client = AsyncOpenAI(
        api_key=settings.openai_api_key, timeout=OPENAI_VISION_TIMEOUT_SECONDS
    )

    images = _pdf_pages_to_base64_images(data)
    if not images:
        return ""

    content: list[ChatCompletionContentPartParam] = [
        {
            "type": "text",
            "text": (
                "Extract ALL text from these scanned syllabus pages. "
                "Return the raw text only, preserving structure like tables and lists. "
                "Do not summarize."
            ),
        },
    ]
    for img_b64 in images:
        content.append(
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{img_b64}",
                    "detail": settings.vision_detail,  # type: ignore[typeddict-item]
                },
            }
        )

    response = await client.chat.completions.create(
        model=settings.vision_model,
        messages=[{"role": "user", "content": content}],
        max_tokens=4096,
    )

    result = response.choices[0].message.content
    return result.strip() if result else ""


async def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF. Tries text extraction first, falls back to vision."""
    try:
        text = await asyncio.wait_for(
            run_in_threadpool(_extract_pdf_text, data),
            timeout=PDF_PROCESSING_TIMEOUT_SECONDS,
        )
    except asyncio.TimeoutError:
        raise HTTPException(
            status_code=408,
            detail="PDF processing timed out. Try a smaller or simpler file.",
        )
    except HTTPException:
        raise
    except fitz.FileDataError:
        raise HTTPException(
            status_code=422, detail="PDF appears to be password-protected or corrupt."
        )

    if len(text.strip()) < 50:
        try:
            vision_text = await _extract_pdf_via_vision(data)
            if vision_text:
                return vision_text
        except Exception:
            raise HTTPException(
                status_code=502,
                detail="OCR failed — try uploading a text-based PDF instead.",
            )

    return text


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        with ZipFile(io.BytesIO(data)) as archive:
            total_uncompressed = sum(info.file_size for info in archive.infolist())
        if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise HTTPException(
                status_code=422,
                detail="DOCX file is too large to process.",
            )
    except BadZipFile:
        raise HTTPException(status_code=422, detail="DOCX file appears to be corrupt.")

    try:
        doc = Document(io.BytesIO(data))
    except BadZipFile:
        raise HTTPException(status_code=422, detail="DOCX file appears to be corrupt.")

    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


async def _extract_images_via_vision(images_data: list[bytes]) -> str:
    """Use OpenAI vision to extract text from one or more screenshot images."""
    client = AsyncOpenAI(
        api_key=settings.openai_api_key, timeout=OPENAI_VISION_TIMEOUT_SECONDS
    )

    content: list[ChatCompletionContentPartParam] = [
        {
            "type": "text",
            "text": (
                "These are screenshots of a course syllabus or assignment page. "
                "Extract ALL text from every image. Return the raw text only, "
                "preserving structure like tables, dates, and lists. "
                "Combine text across images into one coherent document. "
                "Do not summarize."
            ),
        },
    ]
    for img_bytes in images_data:
        b64 = base64.b64encode(img_bytes).decode("utf-8")
        mime = "image/png" if img_bytes[:4] == b"\x89PNG" else "image/jpeg"
        content.append(
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime};base64,{b64}",
                    "detail": settings.vision_detail,  # type: ignore[typeddict-item]
                },
            }
        )

    response = await client.chat.completions.create(
        model=settings.vision_model,
        messages=[{"role": "user", "content": content}],
        max_tokens=4096,
    )

    result = response.choices[0].message.content
    return result.strip() if result else ""


IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp"}

SyncHandler = Callable[[bytes], str]
AsyncHandler = Callable[[bytes], Awaitable[str]]
Handler = SyncHandler | AsyncHandler

CONTENT_TYPE_MAP: dict[str, Handler] = {
    PDF_MIME: _extract_pdf,
    DOCX_MIME: _extract_docx,
}

EXTENSION_MAP: dict[str, Handler] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def _filename_extension(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


async def detect_mime(file: UploadFile) -> str:
    """Sniff the real MIME type from the file's leading bytes."""
    await file.seek(0)
    head = await file.read(8192)
    await file.seek(0)

    if not head:
        raise HTTPException(status_code=400, detail="Empty upload")

    if head.startswith(b"%PDF-"):
        return PDF_MIME
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if head.startswith(b"RIFF") and head[8:12] == b"WEBP":
        return "image/webp"
    if head.startswith(b"PK\x03\x04"):
        if _filename_extension(file.filename) == ".docx":
            return DOCX_MIME
        raise HTTPException(
            status_code=400,
            detail="Unsupported archive file. Upload a PDF, DOCX, or image.",
        )

    kind = filetype.guess(head)
    if kind is not None and kind.mime in ALLOWED_IMAGE_MIMES | ALLOWED_DOC_MIMES:
        return kind.mime

    raise HTTPException(status_code=400, detail="Unrecognized file type")


async def classify_upload(file: UploadFile) -> str:
    """Return 'image' for screenshots or the sniffed MIME for supported documents."""
    mime = await detect_mime(file)
    if mime in ALLOWED_IMAGE_MIMES:
        return "image"
    if mime in ALLOWED_DOC_MIMES:
        return mime
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {mime}")


def _is_image(file: UploadFile) -> bool:
    """Legacy header/extension-based image check. Prefer classify_upload."""
    if file.content_type in IMAGE_CONTENT_TYPES:
        return True
    if file.filename:
        ext = _filename_extension(file.filename)
        return ext in {".png", ".jpg", ".jpeg", ".webp"}
    return False


async def extract_text(file: UploadFile, mime: str | None = None) -> str:
    """Extract plain text from a single uploaded document."""
    if mime is None:
        mime = await detect_mime(file)

    data = await file.read()
    content_hash = extraction_cache.compute_hash(data)

    cached = await extraction_cache.get_cached(content_hash)
    if cached is not None:
        return cached

    handler: Handler | None = CONTENT_TYPE_MAP.get(mime)

    if handler is None and file.filename:
        ext = _filename_extension(file.filename)
        handler = EXTENSION_MAP.get(ext)

    if handler is None:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {mime}",
        )

    if asyncio.iscoroutinefunction(handler):
        text = await handler(data)
    else:
        text = handler(data)  # type: ignore[misc]

    if not text:
        raise HTTPException(
            status_code=422, detail="Could not extract any text from the file."
        )

    assert isinstance(text, str)

    await extraction_cache.put_cached(
        content_hash=content_hash,
        extracted_text=text,
        source_mime=mime,
        vision_model=settings.vision_model,
        vision_used=(mime == PDF_MIME),
        byte_size=len(data),
    )

    return text


async def extract_text_from_images(files: list[UploadFile]) -> str:
    """Extract text from multiple screenshot images via LLM vision."""
    images_data: list[bytes] = []
    for f in files:
        data = await f.read()
        if not data:
            continue
        images_data.append(data)

    if not images_data:
        raise HTTPException(status_code=400, detail="No valid images provided.")

    if len(images_data) > MAX_SCREENSHOT_IMAGES:
        raise HTTPException(
            status_code=400,
            detail=f"Maximum {MAX_SCREENSHOT_IMAGES} screenshots per request.",
        )

    content_hash = extraction_cache.compute_hash_multi(images_data)
    cached = await extraction_cache.get_cached(content_hash)
    if cached is not None:
        return cached

    text = await _extract_images_via_vision(images_data)
    if not text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract any text from the screenshots.",
        )

    await extraction_cache.put_cached(
        content_hash=content_hash,
        extracted_text=text,
        source_mime="image-set",
        vision_model=settings.vision_model,
        vision_used=True,
        byte_size=sum(len(b) for b in images_data),
    )

    return text
